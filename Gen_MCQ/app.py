import os
import time
from flask import Flask, render_template, request, send_file
import pdfplumber
import docx
from docx import Document
from werkzeug.utils import secure_filename
import google.generativeai as genai
from fpdf import FPDF


# -------------------------------
# API Key Configuration
# -------------------------------
os.environ["GOOGLE_API_KEY"] = "AIzaSyABl0ebebrhsP0dYjT-qOec06kYSMnjTvE"
genai.configure(api_key=os.environ["GOOGLE_API_KEY"])
model_mcq = genai.GenerativeModel("models/gemini-2.5-flash")
model_summary = genai.GenerativeModel("models/gemini-2.5-flash")  # For summary


# -------------------------------
# Flask Setup
# -------------------------------
app = Flask(__name__)

UPLOAD_FOLDER = os.path.join(os.getcwd(), 'uploads')
RESULTS_FOLDER = os.path.join(os.getcwd(), 'results')

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(RESULTS_FOLDER, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['RESULTS_FOLDER'] = RESULTS_FOLDER
app.config['ALLOWED_EXTENSIONS'] = {'pdf', 'txt', 'docx'}


# -------------------------------
# Utilities
# -------------------------------
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']


def extract_text_from_file(file_path):
    ext = file_path.rsplit('.', 1)[1].lower()
    if ext == 'pdf':
        with pdfplumber.open(file_path) as pdf:
            text = ''.join([page.extract_text() or '' for page in pdf.pages])
        return text
    elif ext == 'docx':
        doc = docx.Document(file_path)
        return ' '.join([para.text for para in doc.paragraphs])
    elif ext == 'txt':
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    return None


# -------------------------------
# MCQ and Summary Generation
# -------------------------------
def generate_mcqs_with_options(input_text, num_questions, difficulty='Medium', include_answers=True):
    answer_text = "Include the correct answer at the end of each question." if include_answers else "Do not include answers."
    prompt = f"""
You are an AI assistant helping the user generate multiple-choice questions (MCQs) based on the following text:
'{input_text}'
Please generate exactly {num_questions} MCQs of difficulty '{difficulty}'.
{answer_text}
Each question should have:
- A clear question
- Four answer options (A, B, C, D)
Format:
## MCQ
Question: [question]
A) [option A]
B) [option B]
C) [option C]
D) [option D]
Answer: [correct answer]
"""
    response = model_mcq.generate_content(prompt).text.strip()
    return response


def generate_summary(input_text):
    prompt = f"""
Summarize the following text in a single, brief paragraph.
Do NOT use bullets, stars, lists, or numbering.
Keep the summary clear, concise, and easy to read:

{input_text[:12000]}
"""
    response = model_summary.generate_content(prompt).text.strip()
    return response


def save_text_to_file(content, filename):
    path = os.path.join(app.config['RESULTS_FOLDER'], filename)
    with open(path, 'w', encoding='utf-8') as f:
        f.write(content)
    return path


def create_docx(content, filename):
    doc = Document()
    doc.add_heading('Generated MCQs', 0)
    
    # Split content into paragraphs/lines for better formatting
    lines = content.split('\n')
    for line in lines:
        if line.strip():
            if line.strip().startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.', '10.')):
                # Bold numbered questions
                p = doc.add_paragraph()
                run = p.add_run(line.strip())
                run.bold = True
            elif line.strip().startswith(('-', 'A)', 'B)', 'C)', 'D)', 'Answer:')):
                # Italicize options and answers
                p = doc.add_paragraph()
                run = p.add_run(line.strip())
                run.italic = True
            elif line.strip() == '---SUMMARY---':
                doc.add_heading('Summary', level=1)
            else:
                doc.add_paragraph(line.strip())
    
    docx_path = os.path.join(app.config['RESULTS_FOLDER'], filename)
    doc.save(docx_path)
    return docx_path


def sanitize_text(text):
    replacements = {
        '–': '-', '—': '-', '“': '"', '”': '"',
        '‘': "'", '’': "'", '…': '...', '©': '(c)', '®': '(R)'
    }
    for k, v in replacements.items():
        text = text.replace(k, v)
    return text


def create_pdf(content, filename):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    content = sanitize_text(content)
    pdf.multi_cell(0, 10, content)
    pdf_path = os.path.join(app.config['RESULTS_FOLDER'], filename)
    pdf.output(pdf_path)
    return pdf_path


# -------------------------------
# Routes
# -------------------------------
@app.route('/')
def index():
    return render_template('index.html', warning=None)


@app.route('/generate', methods=['POST'])
def generate():
    if 'file' not in request.files:
        return render_template('index.html', warning="Please upload a file.")

    file = request.files['file']

    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)

        text = extract_text_from_file(file_path)
        if not text or not text.strip():
            return render_template('index.html', warning="The uploaded document does not contain enough readable text.")

        requested_questions = int(request.form['num_questions'])
        
        # Max questions limit
        if requested_questions > 100:
            requested_questions = 100

        difficulty = request.form.get('difficulty', 'Medium')
        include_answers = request.form.get('with_answer') == 'on'
        summary_option = request.form.get('summary_based') == 'on'

        word_count = len(text.split())
        possible_questions = max(word_count // 15, 1)

        if requested_questions > possible_questions:
            warning_message = (
                f"⚠ You requested {requested_questions} questions. "
                f"Based on the uploaded document, a maximum of {possible_questions} questions can be generated. "
                "Please reduce the number of questions or upload a longer document."
            )
            return render_template('index.html', warning=warning_message)

        # ---------------------------------
        # Timer Start (Measure MCQ Generation)
        # ---------------------------------
        start_time = time.time()
        mcqs_raw = generate_mcqs_with_options(text, requested_questions, difficulty, include_answers=True)
        end_time = time.time()
        time_taken = round(end_time - start_time, 2)  # In seconds, 2 decimal places

        # -- MCQ Postprocessing --
        mcq_list = [q.strip() for q in mcqs_raw.split("## MCQ") if q.strip()]
        filtered_mcq_list = []
        for block in mcq_list:
            lines = block.splitlines()
            if lines and lines[0].strip().lower().startswith("question:"):
                filtered_mcq_list.append(block)
        if not filtered_mcq_list:
            filtered_mcq_list = mcq_list
        filtered_mcq_list = filtered_mcq_list[:requested_questions]

        if include_answers:
            mcqs = ""
            for idx, mcq in enumerate(filtered_mcq_list, 1):
                mcqs += f"{idx}. {mcq}\n\n"
            key_pdf_filename = None
        else:
            mcqs_no_answers_list = []
            key_answers_list = []

            for idx, mcq in enumerate(filtered_mcq_list, 1):
                parts = mcq.split('\n')
                question_line = parts[0]
                option_lines = []
                answer_line = ""
                for line in parts[1:]:
                    if line.strip().lower().startswith("answer:"):
                        answer_line = line
                    else:
                        option_lines.append(line)
                mcqs_no_answers_list.append(f"{idx}. {question_line}\n" + "\n".join(option_lines))
                if answer_line:
                    key_ans = answer_line.split(":", 1)[1].strip()
                else:
                    key_ans = "(Answer not found)"
                key_answers_list.append(f"{idx}. {key_ans}")

            mcqs = "\n\n".join(mcqs_no_answers_list)
            key_answers_text = "\n".join(key_answers_list)

            # Create PDF for answer key
            key_pdf_filename = f"key_answers_{filename.rsplit('.', 1)[0]}.pdf"
            create_pdf(key_answers_text, key_pdf_filename)

        summary_text = generate_summary(text) if summary_option else None

        content_for_file = mcqs
        if summary_text:
            content_for_file += "\n\n---SUMMARY---\n" + summary_text

        docx_filename = pdf_filename = None
        if content_for_file.strip():
            docx_filename = f"generated_mcqs_{filename.rsplit('.', 1)[0]}.docx"
            pdf_filename = f"generated_mcqs_{filename.rsplit('.', 1)[0]}.pdf"
            create_docx(content_for_file, docx_filename)
            create_pdf(content_for_file, pdf_filename)

        return render_template(
            'results.html',
            mcqs=mcqs,
            summary_text=summary_text,
            docx_filename=docx_filename,  # Changed from txt_filename
            pdf_filename=pdf_filename,
            key_pdf_filename=key_pdf_filename,
            time_taken=time_taken,  # Pass the time taken
            warning=None
        )

    return render_template('index.html', warning="Invalid file format. Please upload PDF, DOCX, or TXT.")


@app.route('/download/<filename>')
def download_file(filename):
    file_path = os.path.join(app.config['RESULTS_FOLDER'], filename)
    if not os.path.exists(file_path):
        return "File not found", 404
    return send_file(file_path, as_attachment=True)


# -------------------------------
# Run App
# -------------------------------
if __name__ == "__main__":
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    os.makedirs(RESULTS_FOLDER, exist_ok=True)
    app.run(debug=True)
